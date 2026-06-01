from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)


OUT = "docs/CV.docx"
PDF_OUT = "docs/CV.pdf"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4D78")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_run(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    set_run(r, 9.5, True, "1F4D78")
    return p


def add_role(doc, title, org, date, location=None):
    p = doc.add_paragraph()
    p.style = doc.styles["Role"]
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run(r, 9.2, True, "000000")
    r = p.add_run(f" | {org}")
    set_run(r, 9.2, False, "000000")
    if location:
        r = p.add_run(f" | {location}")
        set_run(r, 9.2, False, "555555")
    if date:
        r = p.add_run(f" | {date}")
        set_run(r, 9.2, False, "555555")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together = True
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2.2)
    p.paragraph_format.line_spacing = 1.02
    r = p.add_run(text)
    set_run(r, 8.25, None, "000000")


def add_labeled_bullet(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together = True
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2.2)
    p.paragraph_format.line_spacing = 1.02
    r = p.add_run(f"{label}: ")
    set_run(r, 8.25, True, "000000")
    r = p.add_run(text)
    set_run(r, 8.25, None, "000000")


def add_skill(doc, label, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_after = Pt(1.8)
    r = p.add_run(f"{label}: ")
    set_run(r, 8.45, True, "000000")
    r = p.add_run(text)
    set_run(r, 8.45, None, "000000")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(8.6)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.04

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    h1.font.size = Pt(9.5)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string("1F4D78")
    h1.paragraph_format.space_before = Pt(7)
    h1.paragraph_format.space_after = Pt(3)

    role = doc.styles.add_style("Role", 1)
    role.base_style = normal
    role.font.name = "Arial"
    role._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    role.font.size = Pt(9.2)
    role.paragraph_format.space_before = Pt(3)
    role.paragraph_format.space_after = Pt(1)

    for style_name in ("List Bullet",):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(8.25)
        style.paragraph_format.space_after = Pt(2.2)
        style.paragraph_format.line_spacing = 1.02


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.47)
section.bottom_margin = Inches(0.45)
section.left_margin = Inches(0.56)
section.right_margin = Inches(0.56)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)
configure_styles(doc)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(1)
r = title.add_run("MAHSHID ALINOORI")
set_run(r, 18, True, "0B2545")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(3)
r = subtitle.add_run("Machine Learning Engineer")
set_run(r, 10.2, True, "1F4D78")

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(6)
for idx, item in enumerate(
    [
        ("Toronto, ON, Canada", None),
        ("Email", "mailto:mahshid.alinoori@gmail.com"),
        ("LinkedIn", "https://www.linkedin.com/in/mahshid-alinoori/"),
        ("GitHub", "https://github.com/mahshidaln"),
        ("Google Scholar", "https://scholar.google.com/citations?user=4fKZfJkAAAAJ"),
        ("Website", "https://mahshidaln.github.io/"),
    ]
):
    if idx:
        run = contact.add_run(" | ")
        set_run(run, 8.2, None, "555555")
    text, url = item
    if url:
        add_hyperlink(contact, text, url)
    else:
        run = contact.add_run(text)
        set_run(run, 8.2, None, "555555")

summary = doc.add_paragraph()
summary.paragraph_format.space_after = Pt(5)
summary.paragraph_format.line_spacing = 1.04
r = summary.add_run(
    "Machine Learning Engineer with 4+ years of experience turning applied AI research into production systems, "
    "evaluation workflows, and reusable ML infrastructure. Strengths include LLM applications, agentic AI evaluation, "
    "foundation models, cloud deployment, synthetic data generation, and privacy-preserving machine learning."
)
set_run(r, 8.6, None, "000000")

add_heading(doc, "Core Strengths")
for label, text in [
    ("Production AI", "Led LLM-powered analytics and reporting systems from architecture through secure cloud deployment and stakeholder adoption."),
    ("Evaluation & Governance", "Built benchmarking workflows for agentic AI reliability, observability, explainability, and failure analysis."),
    ("Research Engineering", "Implemented training, data processing, and evaluation pipelines for foundation models and applied ML research."),
    ("Technical Enablement", "Designed reproducible bootcamps and reference implementations adopted by 30-40 industry teams per cohort."),
]:
    add_labeled_bullet(doc, label, text)

add_heading(doc, "Skills")
add_skill(doc, "Programming", "Python, SQL, Java, JavaScript, C/C++")
add_skill(
    doc,
    "AI / Machine Learning",
    "PyTorch, Hugging Face Transformers, LLMs, Generative AI, Retrieval-Augmented Generation (RAG), Agentic AI, Foundation Models, Deep Learning, Self-Supervised Learning, Recommender Systems, Multimodal AI, Synthetic Data Generation, Differential Privacy",
)
add_skill(
    doc,
    "LLM & Evaluation",
    "OpenAI, Azure OpenAI, CrewAI, DeepEval, Opik, AI Evaluation, Agent Evaluation, Benchmark Design, Explainable AI (XAI)",
)
add_skill(
    doc,
    "MLOps & Infrastructure",
    "Docker, Kubernetes, Azure, Terraform, CI/CD, GitHub Actions, MLflow, Weights & Biases, vLLM, Linux",
)
add_skill(doc, "Backend & Data", "FastAPI, REST APIs, PostgreSQL, MongoDB, Pinecone, ChromaDB")

add_heading(doc, "Experience")
add_role(
    doc,
    "Applied ML Specialist (Intern to Associate to Applied ML Specialist)",
    "Vector Institute",
    "Jan 2022 - Present",
    "Toronto, Canada",
)
for label, text in [
    (
        "LLM platform delivery",
        "Led a thematic analysis and reporting platform for a major mental health organization, including Azure infrastructure, secure deployment, automated reports, and Tableau analytics.",
    ),
    (
        "Agent evaluation",
        "Built benchmarking workflows with OpenAI, Azure OpenAI, CrewAI, DeepEval, and Opik to assess reliability, observability, explainability, and failure modes.",
    ),
    (
        "Foundation models",
        "Co-developed EHRMamba for long-context electronic health records, contributing data pipelines, training infrastructure, experiments, and evaluation systems.",
    ),
    (
        "Open-source ML",
        "Contributed to CyclOps with healthcare ML data pipelines, evaluation workflows, experimentation tools, and applied ML examples.",
    ),
    (
        "Technical enablement",
        "Delivered advanced bootcamps and reproducible implementations for RAG, multimodal AI, recommender systems, synthetic data, differential privacy, and foundation models.",
    ),
    (
        "Challenge operations",
        "Supported the MIDST Challenge at SaTML 2025 across design, evaluation methodology, participant review, and operations, attracting 700+ submissions.",
    ),
]:
    add_labeled_bullet(doc, label, text)

add_role(doc, "Co-Founder & Chief Innovation Officer", "Echolair", "Apr 2023 - Jan 2024")
for label, text in [
    (
        "Product strategy",
        "Co-founded a generative audio startup for music creators and led technical strategy, AI system design, MVP development, and research-to-product translation.",
    ),
    (
        "Market validation",
        "Led customer discovery, partner conversations, accelerator work, roadmap planning, and product-market fit exploration.",
    ),
]:
    add_labeled_bullet(doc, label, text)

add_role(doc, "Backend Developer", "Papion Software Group", "Jun 2017 - Mar 2018")
for label, text in [
    (
        "Backend services",
        "Developed REST APIs and authentication/object-storage microservices for a social platform serving more than one million users.",
    ),
]:
    add_labeled_bullet(doc, label, text)

add_heading(doc, "Selected Publications")
for item in [
    "Music-STAR: A Style Translation System for Audio-based Re-instrumentation (First Author), ISMIR 2022.",
    "EHRMamba: Towards Generalizable and Scalable Foundation Models for Electronic Health Records, ML4H @ NeurIPS 2024.",
    "From Features to Actions: Explainability in Traditional and Agentic AI Systems.",
    "Transparency in Agentic AI: A Survey of Interpretability, Explainability, and Governance.",
]:
    add_bullet(doc, item)

add_heading(doc, "Education")
add_role(doc, "M.Sc. Computer Science", "York University", "Toronto, Canada")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Thesis: Music-STAR: Audio-Based Multi-Instrument Music Style Translation.")
set_run(r, 8.4, None, "000000")
add_role(doc, "B.Sc. Computer Engineering", "Amirkabir University of Technology", "")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Graduated ranked 1st among hardware engineering students.")
set_run(r, 8.4, None, "000000")

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(2)
    run = footer.add_run("Mahshid Alinoori | CV")
    set_run(run, 7.5, None, "777777")

doc.save(OUT)


def pdf_para(text, style):
    return Paragraph(text.replace("&", "&amp;"), style)


def pdf_bullets(items, style):
    return ListFlowable(
        [ListItem(pdf_para(item, style), leftIndent=12) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=16,
        bulletFontName="Helvetica",
        bulletFontSize=5.8,
        bulletOffsetY=1,
    )


styles = getSampleStyleSheet()
body = ParagraphStyle(
    "Body",
    parent=styles["BodyText"],
    fontName="Helvetica",
    fontSize=10.0,
    leading=12.4,
    spaceAfter=4.0,
)
role_style = ParagraphStyle(
    "Role",
    parent=body,
    fontName="Helvetica",
    fontSize=10.8,
    leading=13.0,
    spaceBefore=7.0,
    spaceAfter=3.0,
)
section_style = ParagraphStyle(
    "Section",
    parent=body,
    fontName="Helvetica-Bold",
    fontSize=12.0,
    leading=14.0,
    textColor=colors.HexColor("#1F4D78"),
    spaceBefore=10.0,
    spaceAfter=4.0,
)
skill_style = ParagraphStyle(
    "Skill",
    parent=body,
    fontSize=9.8,
    leading=12.0,
    spaceAfter=3.0,
)
title_style = ParagraphStyle(
    "Title",
    parent=styles["Title"],
    fontName="Helvetica-Bold",
    fontSize=22,
    leading=25,
    alignment=TA_CENTER,
    textColor=colors.HexColor("#0B2545"),
    spaceAfter=0,
)
subtitle_style = ParagraphStyle(
    "Subtitle",
    parent=body,
    fontName="Helvetica-Bold",
    fontSize=12.0,
    leading=14,
    alignment=TA_CENTER,
    textColor=colors.HexColor("#1F4D78"),
    spaceAfter=1,
)
contact_style = ParagraphStyle(
    "Contact",
    parent=body,
    fontSize=9.6,
    leading=11.5,
    alignment=TA_CENTER,
    textColor=colors.HexColor("#555555"),
    spaceAfter=5.0,
)


summary_text = (
    "Machine Learning Engineer with 4+ years of experience turning applied AI research into production systems, "
    "evaluation workflows, and reusable ML infrastructure. Strengths include LLM applications, agentic AI evaluation, "
    "foundation models, cloud deployment, synthetic data generation, and privacy-preserving machine learning."
)
core_strengths = [
    ("Production AI", "Led LLM analytics and reporting systems from architecture through secure deployment and adoption."),
    ("Evaluation & Governance", "Built benchmarks for agent reliability, observability, explainability, and failure analysis."),
    ("Research Engineering", "Implemented training, data processing, and evaluation pipelines for foundation models."),
    ("Technical Enablement", "Created bootcamps and reference implementations for 30-40 industry teams per cohort."),
]
skills = [
    ("Programming", "Python, SQL, Java, JavaScript, C/C++"),
    (
        "AI / Machine Learning",
        "PyTorch, Hugging Face Transformers, LLMs, Generative AI, Retrieval-Augmented Generation (RAG), Agentic AI, Foundation Models, Deep Learning, Self-Supervised Learning, Recommender Systems, Multimodal AI, Synthetic Data Generation, Differential Privacy",
    ),
    (
        "LLM & Evaluation",
        "OpenAI, Azure OpenAI, CrewAI, DeepEval, Opik, AI Evaluation, Agent Evaluation, Benchmark Design, Explainable AI (XAI)",
    ),
    (
        "MLOps & Infrastructure",
        "Docker, Kubernetes, Azure, Terraform, CI/CD, GitHub Actions, MLflow, Weights & Biases, vLLM, Linux",
    ),
    ("Backend & Data", "FastAPI, REST APIs, PostgreSQL, MongoDB, Pinecone, ChromaDB"),
]
vector_bullets = [
    ("LLM platform delivery", "Led a thematic analysis and reporting platform for a major mental health organization, including Azure infrastructure, secure deployment, automated reports, and Tableau analytics."),
    ("Agent evaluation", "Built benchmarking workflows with OpenAI, Azure OpenAI, CrewAI, DeepEval, and Opik to assess reliability, observability, explainability, and failure modes."),
    ("Foundation models", "Co-developed EHRMamba for long-context electronic health records, contributing data pipelines, training infrastructure, experiments, and evaluation systems."),
    ("Open-source ML", "Contributed to CyclOps with healthcare ML data pipelines, evaluation workflows, experimentation tools, and applied ML examples."),
    ("Technical enablement", "Delivered advanced bootcamps and reproducible implementations for RAG, multimodal AI, recommender systems, synthetic data, differential privacy, and foundation models."),
    ("Challenge operations", "Supported the MIDST Challenge at SaTML 2025 across design, evaluation methodology, participant review, and operations, attracting 700+ submissions."),
]
echolair_bullets = [
    ("Product strategy", "Co-founded a generative audio startup for music creators and led technical strategy, AI system design, MVP development, and research-to-product translation."),
    ("Market validation", "Led customer discovery, partner conversations, accelerator work, roadmap planning, and product-market fit exploration."),
]
papion_bullets = [
    ("Backend services", "Developed REST APIs and authentication/object-storage microservices for a social platform serving more than one million users."),
]
publications = [
    "Music-STAR: A Style Translation System for Audio-based Re-instrumentation (First Author), ISMIR 2022.",
    "EHRMamba: Towards Generalizable and Scalable Foundation Models for Electronic Health Records, ML4H @ NeurIPS 2024.",
    "From Features to Actions: Explainability in Traditional and Agentic AI Systems.",
    "Transparency in Agentic AI: A Survey of Interpretability, Explainability, and Governance.",
]

pdf = SimpleDocTemplate(
    PDF_OUT,
    pagesize=letter,
    leftMargin=0.56 * inch,
    rightMargin=0.56 * inch,
    topMargin=0.5 * inch,
    bottomMargin=0.5 * inch,
)
story = [
    Paragraph("MAHSHID ALINOORI", title_style),
    Paragraph("Machine Learning Engineer", subtitle_style),
    Paragraph(
        'Toronto, ON, Canada | <a href="mailto:mahshid.alinoori@gmail.com">Email</a> | '
        '<a href="https://www.linkedin.com/in/mahshid-alinoori/">LinkedIn</a> | '
        '<a href="https://github.com/mahshidaln">GitHub</a> | '
        '<a href="https://scholar.google.com/citations?user=4fKZfJkAAAAJ">Google Scholar</a> | '
        '<a href="https://mahshidaln.github.io/">Website</a>',
        contact_style,
    ),
    HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#DADCE0"), spaceAfter=5.0),
    pdf_para(summary_text, body),
    Paragraph("CORE STRENGTHS", section_style),
    pdf_bullets([f"<b>{label}:</b> {text}" for label, text in core_strengths], body),
    Paragraph("SKILLS", section_style),
]
for label, text in skills:
    story.append(Paragraph(f"<b>{label}:</b> {text}", skill_style))
story.extend(
    [
        Paragraph("EXPERIENCE", section_style),
        Paragraph(
            "<b>Applied ML Specialist (Intern to Associate to Applied ML Specialist)</b> | Vector Institute | Toronto, Canada | Jan 2022 - Present",
            role_style,
        ),
        pdf_bullets([f"<b>{label}:</b> {text}" for label, text in vector_bullets], body),
        Paragraph("<b>Co-Founder & Chief Innovation Officer</b> | Echolair | Apr 2023 - Jan 2024", role_style),
        pdf_bullets([f"<b>{label}:</b> {text}" for label, text in echolair_bullets], body),
        Paragraph("<b>Backend Developer</b> | Papion Software Group | Jun 2017 - Mar 2018", role_style),
        pdf_bullets([f"<b>{label}:</b> {text}" for label, text in papion_bullets], body),
        Paragraph("SELECTED PUBLICATIONS", section_style),
        pdf_bullets(publications, body),
        Paragraph("EDUCATION", section_style),
        Paragraph("<b>M.Sc. Computer Science</b> | York University | Toronto, Canada", role_style),
        pdf_para("Thesis: Music-STAR: Audio-Based Multi-Instrument Music Style Translation.", body),
        Paragraph("<b>B.Sc. Computer Engineering</b> | Amirkabir University of Technology", role_style),
        pdf_para("Graduated ranked 1st among hardware engineering students.", body),
    ]
)

pdf.build(story)
print(OUT)
print(PDF_OUT)
